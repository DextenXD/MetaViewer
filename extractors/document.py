from docx import Document
import zipfile
import json

def extract_doc(filepath):
  try:
    doc = Document(filepath)
    props = doc.core_properties
    data = {
      # Core props
      "Author": props.author or "Unknown",
      "Last Modified By": props.last_modified_by or "Unknown",
      "Created": str(props.created) or "Unknown",
      "Modified": str(props.modified) or "Unknown",
      "Title": props.title or "Unknown",
      "Subject": props.subject or "Unknown",
      "Keywords": props.keywords or "Unknown",
      "Category": props.category or "Unknown",
      "Revision": str(props.revision) or "Unknown",
      "Version": props.version or "Unknown",
      "Content Status": props.content_status or "Unknown",
      "Language": props.language or "Unknown",
      "Identifier": props.identifier or "Unknown",

      # Docu stats
      "Paragraphs": len(doc.paragraphs),
      "Tables": len(doc.tables),
      "Sections": len(doc.sections),
      "Images": len(doc.inline_shapes),
    }

    return data
  except FileNotFoundError:
    print(f"File not found {filepath}")
    return None
  except Exception as e:
    print(f"How did you even manage to do this??? {e}")
    return None